from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .exceptions import IndicatorParsingError, IndicatorValidationError


COMPETENCE_CODE_RE = re.compile(
    r'(?<![\w])(?P<code>(?:УК|ОПК|ПК)\s*[-–—]\s*\d+)(?!\s*\.\s*\d)',
    re.IGNORECASE,
)
INDICATOR_CODE_RE = re.compile(
    r'(?<![\w])(?P<code>(?:УК|ОПК|ПК)\s*[-–—]\s*\d+\s*\.\s*\d+)(?:\s*[.:])?\s*',
    re.IGNORECASE,
)


def normalize_text(value: str | None) -> str:
    if value is None:
        return ''
    return ' '.join(str(value).replace('\u00a0', ' ').strip().split())


def normalize_code(value: str | None) -> str:
    text = normalize_text(value).upper()
    text = re.sub(r'\s*[-–—]\s*', '-', text)
    text = re.sub(r'\s*\.\s*', '.', text)
    return text.rstrip('.')


@dataclass(frozen=True, slots=True)
class ParsedIndicatorRow:
    competence_code: str
    indicator_code: str
    text: str
    source_table_number: int
    source_row_number: int


@dataclass(frozen=True, slots=True)
class ParsedIndicatorDocument:
    source_filename: str
    tables_found: int
    skipped_rows: int
    rows: list[ParsedIndicatorRow]


class IndicatorDocxParser:
    """Извлекает строки индикаторов из смысловых таблиц DOCX."""

    def parse_upload(self, uploaded_file) -> ParsedIndicatorDocument:
        from .indicator_document_converter import IndicatorDocumentConverter

        prepared = IndicatorDocumentConverter().prepare_upload(uploaded_file)
        return self.parse_bytes(
            content=prepared.docx_content,
            source_filename=prepared.source_filename,
        )

    def parse_bytes(self, *, content: bytes, source_filename: str) -> ParsedIndicatorDocument:
        if not content:
            raise IndicatorValidationError('Загружен пустой файл.')
        try:
            document = Document(BytesIO(content))
        except (PackageNotFoundError, BadZipFile, ValueError, KeyError) as exc:
            raise IndicatorParsingError('Файл повреждён или не является корректным DOCX-документом.') from exc
        except Exception as exc:
            raise IndicatorParsingError('Не удалось прочитать структуру DOCX-документа.') from exc

        rows: list[ParsedIndicatorRow] = []
        tables_found = 0
        skipped_rows = 0
        for table_number, table in enumerate(document.tables, start=1):
            header = self._find_header(table)
            if header is None:
                continue
            tables_found += 1
            header_row_index, competence_column, indicator_column = header
            for row_index, row in enumerate(table.rows[header_row_index + 1 :], start=header_row_index + 2):
                competence_text = normalize_text(row.cells[competence_column].text)
                indicator_text = normalize_text(row.cells[indicator_column].text)
                if not competence_text and not indicator_text:
                    skipped_rows += 1
                    continue

                competence_code = self._extract_competence_code(competence_text)
                if (
                    not competence_code
                    and competence_text
                    and competence_text.casefold() == indicator_text.casefold()
                    and not INDICATOR_CODE_RE.search(indicator_text)
                ):
                    skipped_rows += 1
                    continue
                parsed = self._split_indicators(
                    competence_code=competence_code,
                    indicator_text=indicator_text,
                    table_number=table_number,
                    row_number=row_index,
                )
                if parsed:
                    rows.extend(parsed)
                else:
                    rows.append(
                        ParsedIndicatorRow(
                            competence_code=competence_code,
                            indicator_code='',
                            text=indicator_text,
                            source_table_number=table_number,
                            source_row_number=row_index,
                        )
                    )

        if not tables_found:
            raise IndicatorParsingError(
                'В документе не найдена таблица с колонкой «Индикаторы достижения компетенции».'
            )
        if not rows:
            raise IndicatorParsingError('В таблицах индикаторов не найдено строк для импорта.')

        return ParsedIndicatorDocument(
            source_filename=source_filename,
            tables_found=tables_found,
            skipped_rows=skipped_rows,
            rows=rows,
        )

    @staticmethod
    def _find_header(table) -> tuple[int, int, int] | None:
        for row_index, row in enumerate(table.rows[:5]):
            cells = [normalize_text(cell.text).casefold() for cell in row.cells]
            indicator_columns = [
                index
                for index, text in enumerate(cells)
                if 'индикатор' in text and 'достижен' in text and 'компетенц' in text
            ]
            competence_columns = [
                index
                for index, text in enumerate(cells)
                if 'компетенц' in text and 'код' in text and index not in indicator_columns
            ]
            if indicator_columns and competence_columns:
                return row_index, competence_columns[0], indicator_columns[0]
        return None

    @staticmethod
    def _extract_competence_code(value: str) -> str:
        match = COMPETENCE_CODE_RE.search(value)
        return normalize_code(match.group('code')) if match else ''

    @staticmethod
    def _split_indicators(
        *,
        competence_code: str,
        indicator_text: str,
        table_number: int,
        row_number: int,
    ) -> list[ParsedIndicatorRow]:
        matches = list(INDICATOR_CODE_RE.finditer(indicator_text))
        result = []
        for index, match in enumerate(matches):
            text_start = match.end()
            text_end = matches[index + 1].start() if index + 1 < len(matches) else len(indicator_text)
            result.append(
                ParsedIndicatorRow(
                    competence_code=competence_code,
                    indicator_code=normalize_code(match.group('code')),
                    text=normalize_text(indicator_text[text_start:text_end]).strip(' .;'),
                    source_table_number=table_number,
                    source_row_number=row_number,
                )
            )
        return result
