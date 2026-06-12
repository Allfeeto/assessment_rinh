from __future__ import annotations

import os
import json
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from secrets import token_hex

from docx import Document

from .exceptions import IndicatorParsingError, IndicatorValidationError
from .indicator_parser import normalize_text


MAX_INDICATOR_WORD_SIZE = 10 * 1024 * 1024
SUPPORTED_WORD_EXTENSIONS = {'.doc', '.docx'}


@dataclass(frozen=True, slots=True)
class PreparedIndicatorDocument:
    source_filename: str
    original_content: bytes
    docx_content: bytes


class IndicatorDocumentConverter:
    """Подготавливает старый бинарный DOC или DOCX для общего DOCX-парсера."""

    def prepare_upload(self, uploaded_file) -> PreparedIndicatorDocument:
        filename, content = self.read_upload(uploaded_file)
        return self.prepare_content(source_filename=filename, content=content)

    @staticmethod
    def read_upload(uploaded_file) -> tuple[str, bytes]:
        if uploaded_file is None:
            raise IndicatorValidationError('Файл не выбран.')

        filename = normalize_text(getattr(uploaded_file, 'name', ''))
        if not filename:
            raise IndicatorValidationError('Не удалось определить имя загруженного файла.')
        if len(filename) > 255:
            raise IndicatorValidationError('Имя файла Word не должно превышать 255 символов.')
        extension = Path(filename).suffix.casefold()
        if extension not in SUPPORTED_WORD_EXTENSIONS:
            raise IndicatorValidationError('Поддерживаются только файлы Word с расширением .doc или .docx.')

        try:
            content = uploaded_file.read()
        except Exception as exc:
            raise IndicatorValidationError('Не удалось прочитать загруженный файл.') from exc
        if not content:
            raise IndicatorValidationError('Загружен пустой файл.')
        if len(content) > MAX_INDICATOR_WORD_SIZE:
            raise IndicatorValidationError('Размер файла Word не должен превышать 10 МБ.')
        return filename, content

    def prepare_content(self, *, source_filename: str, content: bytes) -> PreparedIndicatorDocument:
        extension = Path(source_filename).suffix.casefold()
        docx_content = content if extension == '.docx' else self.convert_doc_bytes(
            content=content,
            source_filename=source_filename,
        )
        return PreparedIndicatorDocument(
            source_filename=source_filename,
            original_content=content,
            docx_content=docx_content,
        )

    def convert_doc_bytes(self, *, content: bytes, source_filename: str) -> bytes:
        temp_path = Path(tempfile.gettempdir()) / f'indicator-doc-{token_hex(12)}'
        if os.name == 'nt':
            temp_path.mkdir()
        else:
            temp_path.mkdir(mode=0o700)
        try:
            source_path = temp_path / 'source.doc'
            source_path.write_bytes(content)

            converter = shutil.which('soffice') or shutil.which('libreoffice')
            if converter:
                output_path = source_path.with_suffix('.docx')
                self._convert_with_libreoffice(converter, source_path, temp_path)
            elif os.name == 'nt' and shutil.which('powershell.exe'):
                output_path = temp_path / 'converted.docx'
                self._convert_with_microsoft_word(source_path, output_path)
            else:
                raise IndicatorParsingError(
                    'Для обработки .doc на сервере не найден LibreOffice. '
                    'Установите libreoffice-writer или используйте проектный Docker-образ.'
                )

            if not output_path.is_file() or output_path.stat().st_size == 0:
                raise IndicatorParsingError(
                    f'Не удалось преобразовать файл «{source_filename}» из .doc в .docx.'
                )
            return output_path.read_bytes()
        finally:
            shutil.rmtree(temp_path, ignore_errors=True)

    @staticmethod
    def _convert_with_libreoffice(converter: str, source_path: Path, output_dir: Path) -> None:
        env = os.environ.copy()
        env['HOME'] = str(output_dir)
        try:
            result = subprocess.run(
                [
                    converter,
                    '--headless',
                    '--convert-to',
                    'docx',
                    '--outdir',
                    str(output_dir),
                    str(source_path),
                ],
                check=False,
                capture_output=True,
                text=True,
                timeout=60,
                env=env,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise IndicatorParsingError('Не удалось запустить LibreOffice для обработки .doc.') from exc
        if result.returncode != 0:
            raise IndicatorParsingError('LibreOffice не смог преобразовать загруженный .doc-файл.')

    @staticmethod
    def _convert_with_microsoft_word(source_path: Path, output_path: Path) -> None:
        source_path = source_path.resolve()
        output_path = output_path.resolve()
        script_path = source_path.parent / 'convert-doc.ps1'
        script_path.write_text(
            r"""
param([string]$Source)
[Console]::OutputEncoding = [System.Text.UTF8Encoding]::new($false)
$OutputEncoding = [Console]::OutputEncoding
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$word.AutomationSecurity = 3
$doc = $null
try {
    $doc = $word.Documents.Open($Source, $false, $true)
    $tables = @()
    foreach ($table in $doc.Tables) {
        $tableRows = @()
        for ($rowNumber = 1; $rowNumber -le $table.Rows.Count; $rowNumber++) {
            $cells = @()
            for ($columnNumber = 1; $columnNumber -le $table.Columns.Count; $columnNumber++) {
                try {
                    $value = $table.Cell($rowNumber, $columnNumber).Range.Text
                }
                catch {
                    $value = ''
                }
                $value = ($value -replace '[\r\a]+', ' ' -replace '\s+', ' ').Trim()
                $cells += $value
            }
            $tableRows += [PSCustomObject]@{ cells = $cells }
        }
        $tables += [PSCustomObject]@{ rows = $tableRows }
    }
    ConvertTo-Json -InputObject $tables -Depth 8 -Compress
}
finally {
    if ($null -ne $doc) {
        $doc.Close($false)
        [void][System.Runtime.InteropServices.Marshal]::FinalReleaseComObject($doc)
    }
    $word.Quit()
    [void][System.Runtime.InteropServices.Marshal]::FinalReleaseComObject($word)
    [GC]::Collect()
    [GC]::WaitForPendingFinalizers()
}
""".strip(),
            encoding='utf-8-sig',
        )
        try:
            result = subprocess.run(
                [
                    'powershell.exe',
                    '-NoProfile',
                    '-NonInteractive',
                    '-ExecutionPolicy',
                    'Bypass',
                    '-File',
                    str(script_path.resolve()),
                    str(source_path),
                ],
                check=False,
                capture_output=True,
                timeout=60,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise IndicatorParsingError('Не удалось запустить Microsoft Word для обработки .doc.') from exc
        if result.returncode != 0:
            raise IndicatorParsingError('Microsoft Word не смог прочитать таблицы загруженного .doc-файла.')

        try:
            tables = json.loads(result.stdout.decode('utf-8-sig'))
        except (UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise IndicatorParsingError('Microsoft Word вернул некорректную структуру таблиц .doc-файла.') from exc
        if isinstance(tables, dict):
            tables = [tables]
        if not isinstance(tables, list) or not tables:
            raise IndicatorParsingError('В загруженном .doc-файле не найдены таблицы.')

        document = Document()
        for source_table in tables:
            source_rows = source_table.get('rows') if isinstance(source_table, dict) else None
            if isinstance(source_rows, dict):
                source_rows = [source_rows]
            if not isinstance(source_rows, list) or not source_rows:
                continue
            column_count = max(
                (
                    len(row.get('cells', []))
                    for row in source_rows
                    if isinstance(row, dict)
                ),
                default=0,
            )
            if column_count == 0:
                continue
            target_table = document.add_table(rows=len(source_rows), cols=column_count)
            for row_number, source_row in enumerate(source_rows):
                if not isinstance(source_row, dict):
                    continue
                source_cells = source_row.get('cells', [])
                if not isinstance(source_cells, list):
                    source_cells = [source_cells]
                for column_number, value in enumerate(source_cells[:column_count]):
                    target_table.cell(row_number, column_number).text = str(value or '')
        document.save(output_path)
