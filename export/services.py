from __future__ import annotations

import logging
import random
from datetime import datetime
from io import BytesIO
from pathlib import Path

from django.db.models import Q
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from assessment.models import AssessmentItem
from assessment.services import (
    TYPE_MATCHING,
    TYPE_MULTIPLE,
    TYPE_OPEN,
    TYPE_SEQUENCE,
    TYPE_SINGLE,
    get_item_competences,
    get_item_type_ui_name,
    infer_item_type_code,
)
from disciplines.models import ProgramDiscipline

MAKET_PATH = Path(__file__).resolve().parent.parent / 'templates' / 'export' / 'maket.docx'
FONT_NAME = 'Times New Roman'
FONT_SIZE_PT = 12
TITLE_FONT_SIZE_PT = 14
MAX_EXPORT_ITEMS = 1000
RUS_LETTERS = ['А', 'Б', 'В', 'Г', 'Д', 'Е', 'Ж', 'З', 'И', 'К', 'Л', 'М', 'Н', 'О', 'П', 'Р', 'С', 'Т']
logger = logging.getLogger(__name__)


class WordExportError(ValueError):
    status_code = 400


class WordExportNotFoundError(WordExportError):
    status_code = 404


ANSWER_INSTRUCTION_BY_TYPE = {
    TYPE_MATCHING: 'Запишите выбранные цифры под соответствующими буквами, каждый элемент правого столбца используется один раз:',
    TYPE_SEQUENCE: 'Запишите соответствующую последовательность цифр слева направо:',
    TYPE_MULTIPLE: 'Варианты ответа.',
    TYPE_SINGLE: 'Вариант ответа.',
    TYPE_OPEN: 'Укажите правильный ответ:',
}

TASK_INTRO_BY_TYPE = {
    TYPE_MATCHING: 'Прочитайте текст и установите соответствие.',
    TYPE_SEQUENCE: 'Прочитайте текст и установите последовательность.',
    TYPE_MULTIPLE: 'Прочитайте текст, выберите все правильные варианты ответа.',
    TYPE_SINGLE: 'Прочитайте текст и выберите правильный ответ.',
    TYPE_OPEN: 'Прочитайте текст и запишите ответ.',
}

CRITERIA_BY_TYPE = {
    TYPE_MATCHING: (
        '1 балл – полное совпадение с верным ответом; '
        '0 баллов – неверный ответ или его отсутствие.'
    ),
    TYPE_SEQUENCE: (
        '1 балл – полное совпадение с верным ответом; '
        '0 баллов – если допущены ошибки или ответ отсутствует.'
    ),
    TYPE_MULTIPLE: (
        '1 балл – полное совпадение с верным ответом; '
        '0 баллов – неверный ответ или его отсутствие.'
    ),
    TYPE_SINGLE: (
        '1 балл – полное совпадение с верным ответом; '
        '0 баллов – допущены ошибки или ответ отсутствует.'
    ),
    TYPE_OPEN: (
        '2 балла – полный правильный ответ на задание / полное совпадение с эталоном ответа; '
        '1 балл – допущена одна ошибка / неточность / ответ правильный, но не полный; '
        '0 баллов – допущено более одной ошибки / ответ неправильный / ответ отсутствует.'
    ),
}


def _letter(index: int) -> str:
    if 0 <= index < len(RUS_LETTERS):
        return RUS_LETTERS[index]
    return str(index + 1)


def _filtered_items(program_id, discipline_id, filters):
    queryset = (
        AssessmentItem.objects.select_related(
            'assessment_item_type',
            'competence',
            'program_discipline__discipline',
            'program_discipline__educational_program__program_profile',
        )
        .prefetch_related('rows')
        .filter(
            program_discipline__educational_program_id=program_id,
            program_discipline__discipline_id=discipline_id,
        )
        .order_by('id')
    )

    if filters.get('assessment_item_type_id'):
        queryset = queryset.filter(assessment_item_type_id=filters['assessment_item_type_id'])

    if filters.get('competence_id'):
        queryset = queryset.filter(
            Q(competence_id=filters['competence_id'])
            | Q(competence_links__competence_id=filters['competence_id'])
        )

    return queryset.prefetch_related('competence_links__competence').distinct()


def _cleanup_text(value):
    return (value or '').strip()


def _set_cell_text(cell, text: str, *, bold: bool = False, align=None):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    _normalize_paragraph(paragraph)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    if bold:
        run.bold = True


def _iter_rows(item):
    return list(item.rows.order_by('sort_order', 'id'))


def _resolve_indicators_text(item) -> str:
    # Расширяемая точка: если в модели появятся индикаторы, можно брать здесь.
    return '—'


def _prepare_matching_rows(rows, rng: random.Random):
    pairs = []
    distractors = []

    for row in rows:
        left = _cleanup_text(row.left_text)
        right = _cleanup_text(row.right_text)
        if not right:
            continue
        data = {'id': row.id, 'left': left, 'right': right}
        if left:
            pairs.append(data)
        else:
            distractors.append(data)

    right_items = [{'id': pair['id'], 'text': pair['right']} for pair in pairs]
    right_items.extend({'id': row['id'], 'text': row['right']} for row in distractors)
    rng.shuffle(right_items)
    right_index = {entry['id']: idx + 1 for idx, entry in enumerate(right_items)}

    letters = [_letter(i) for i in range(len(pairs))]
    key_parts = [f'{letters[idx]}-{right_index[pair["id"]]}' for idx, pair in enumerate(pairs)]

    return {
        'pairs': pairs,
        'right_items': right_items,
        'letters': letters,
        'answer_key': ', '.join(key_parts),
    }


def _prepare_sequence_rows(rows, rng: random.Random):
    source_rows = [row for row in rows if _cleanup_text(row.left_text)]
    source_rows.sort(key=lambda row: (row.sort_order or 9999, row.id))

    steps = []
    for index, row in enumerate(source_rows, start=1):
        correct_order = row.correct_order if row.correct_order is not None else index
        steps.append({'id': row.id, 'text': _cleanup_text(row.left_text), 'correct_order': correct_order})

    rng.shuffle(steps)
    visible_index = {step['id']: idx + 1 for idx, step in enumerate(steps)}
    correct_chain = sorted(steps, key=lambda step: step['correct_order'])
    answer_key = ', '.join(str(visible_index[step['id']]) for step in correct_chain)
    return {
        'steps': steps,
        'answer_cells': max(len(steps), 1),
        'answer_key': answer_key,
    }


def _list_item_text(index: int, text: str, is_last: bool) -> str:
    suffix = '.' if is_last else ';'
    normalized = (text or '').strip().rstrip(' ;:,.')
    return f'{index}) {normalized}{suffix}'


def _prepare_choice_rows(rows, rng: random.Random, multiple: bool):
    options = [
        {'id': row.id, 'text': _cleanup_text(row.left_text), 'is_correct': bool(row.is_correct)}
        for row in rows
        if _cleanup_text(row.left_text)
    ]
    rng.shuffle(options)

    correct_positions = [idx + 1 for idx, option in enumerate(options) if option['is_correct']]
    if multiple:
        answer_key = ', '.join(str(position) for position in correct_positions)
    else:
        answer_key = str(correct_positions[0]) if correct_positions else ''

    return {
        'options': options,
        'answer_key': answer_key,
    }


def _prepare_open_rows(rows):
    answers = [_cleanup_text(row.open_answer_text) for row in rows if _cleanup_text(row.open_answer_text)]
    return {
        'answers': answers,
        'answer_key': '; '.join(answers),
    }


def _prepare_export_item(item, number: int, rng: random.Random):
    rows = _iter_rows(item)
    type_name = get_item_type_ui_name(item.assessment_item_type)
    type_code = infer_item_type_code(item.assessment_item_type)

    if type_code not in {TYPE_MATCHING, TYPE_SEQUENCE, TYPE_MULTIPLE, TYPE_SINGLE, TYPE_OPEN}:
        logger.warning(
            'Unsupported assessment item type during Word export',
            extra={
                'assessment_item_id': item.id,
                'assessment_item_type_id': item.assessment_item_type_id,
                'assessment_item_type_name': item.assessment_item_type.name,
            },
        )
        raise WordExportError(
            'В выбранном наборе есть задания с неподдерживаемым типом. '
            'Проверьте типы заданий и повторите экспорт.'
        )

    prepared = {
        'number': number,
        'item': item,
        'type_name': type_name,
        'type_code': type_code,
        'task_intro': TASK_INTRO_BY_TYPE.get(type_code, ''),
        'prompt_text': _cleanup_text(item.prompt_text),
        'criteria': CRITERIA_BY_TYPE.get(type_code, ''),
        'answer_instruction': ANSWER_INSTRUCTION_BY_TYPE.get(type_code, ''),
        'indicator_text': _resolve_indicators_text(item),
        'answer_key': '',
        'payload': {},
        'competence_text': '; '.join(
            f'{competence.code} — {competence.name}'
            for competence in get_item_competences(item)
        ) or '—',
    }

    if type_code == TYPE_MATCHING:
        payload = _prepare_matching_rows(rows, rng)
        prepared['payload'] = payload
        prepared['answer_key'] = payload['answer_key']
    elif type_code == TYPE_SEQUENCE:
        payload = _prepare_sequence_rows(rows, rng)
        prepared['payload'] = payload
        prepared['answer_key'] = payload['answer_key']
    elif type_code == TYPE_MULTIPLE:
        payload = _prepare_choice_rows(rows, rng, multiple=True)
        prepared['payload'] = payload
        prepared['answer_key'] = payload['answer_key']
    elif type_code == TYPE_SINGLE:
        payload = _prepare_choice_rows(rows, rng, multiple=False)
        prepared['payload'] = payload
        prepared['answer_key'] = payload['answer_key']
    elif type_code == TYPE_OPEN:
        payload = _prepare_open_rows(rows)
        prepared['payload'] = payload
        prepared['answer_key'] = payload['answer_key']

    return prepared


def _clear_document_content(doc: Document):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith('sectPr'):
            continue
        body.remove(child)


def _normalize_paragraph(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)


def _configure_document_styles(doc: Document):
    for style_name in ('Normal', 'List Paragraph'):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = Pt(FONT_SIZE_PT)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1


def _set_table_col_widths(table, widths_cm: list[float]):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def _add_text(
    doc: Document,
    text: str = '',
    style: str = 'Normal',
    *,
    align=None,
    size_pt: int = FONT_SIZE_PT,
):
    paragraph = doc.add_paragraph(style=style)
    _normalize_paragraph(paragraph)
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(size_pt)
    return paragraph


def _add_text_styled(
    doc: Document,
    text: str = '',
    style: str = 'Normal',
    *,
    italic: bool = False,
    bold: bool = False,
    align=None,
    size_pt: int = FONT_SIZE_PT,
):
    paragraph = doc.add_paragraph(style=style)
    _normalize_paragraph(paragraph)
    if text:
        run = paragraph.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(size_pt)
        if italic:
            run.italic = True
        if bold:
            run.bold = True
    if align is not None:
        paragraph.alignment = align
    return paragraph


def _add_italic_line(doc: Document, text: str):
    paragraph = doc.add_paragraph(style='Normal')
    _normalize_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    run.italic = True
    return paragraph


def _add_type_line(doc: Document, type_name: str):
    paragraph = doc.add_paragraph(style='Normal')
    _normalize_paragraph(paragraph)
    label = paragraph.add_run('Тип задания: ')
    label.font.name = FONT_NAME
    label.font.size = Pt(FONT_SIZE_PT)
    label.italic = True
    value = paragraph.add_run(type_name)
    value.font.name = FONT_NAME
    value.font.size = Pt(FONT_SIZE_PT)
    return paragraph


def _render_matching(doc: Document, prepared: dict):
    item = prepared['item']
    payload = prepared['payload']

    _add_text(
        doc,
        'К каждой позиции, данной в левом столбце, подберите соответствующую позицию из правого столбца:',
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    _add_text(doc, '')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    left_header = table.rows[0].cells[0]
    _set_cell_text(
        left_header,
        _cleanup_text(item.left_column_title) or 'Левый столбец',
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    left_header.merge(table.rows[0].cells[1])
    right_header = table.rows[0].cells[2]
    _set_cell_text(
        right_header,
        _cleanup_text(item.right_column_title) or 'Правый столбец',
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    right_header.merge(table.rows[0].cells[3])

    right_entries = payload['right_items']
    left_pairs = payload['pairs']
    row_count = max(len(left_pairs), len(right_entries))

    for idx in range(row_count):
        cells = table.add_row().cells
        if idx < len(left_pairs):
            _set_cell_text(cells[0], payload['letters'][idx])
            _set_cell_text(cells[1], left_pairs[idx]['left'])
        if idx < len(right_entries):
            _set_cell_text(cells[2], str(idx + 1))
            _set_cell_text(cells[3], right_entries[idx]['text'])
    _set_table_col_widths(table, [0.8, 6.7, 0.8, 6.7])

    _add_text(doc, '')
    _add_text(
        doc,
        prepared['answer_instruction'],
    )
    _add_text(
        doc,
        '',
    )

    answer_table = doc.add_table(rows=2, cols=max(len(payload['letters']), 1))
    answer_table.style = 'Table Grid'
    _set_table_col_widths(answer_table, [1.0] * max(len(payload['letters']), 1))
    for idx, letter in enumerate(payload['letters'] or ['']):
        _set_cell_text(answer_table.rows[0].cells[idx], letter)
    for idx in range(max(len(payload['letters']), 1)):
        _set_cell_text(answer_table.rows[1].cells[idx], '')
    _add_text(doc, '')


def _render_sequence(doc: Document, prepared: dict):
    payload = prepared['payload']
    options_count = len(payload['steps'])
    for idx, step in enumerate(payload['steps'], start=1):
        _add_text(doc, _list_item_text(idx, step['text'], idx == options_count))
    _add_text(doc, '')
    _add_text(doc, prepared['answer_instruction'])
    _add_text(doc, '')
    answer_table = doc.add_table(rows=1, cols=payload['answer_cells'])
    answer_table.style = 'Table Grid'
    _set_table_col_widths(answer_table, [1.0] * payload['answer_cells'])
    for idx in range(payload['answer_cells']):
        _set_cell_text(answer_table.rows[0].cells[idx], '')
    _add_text(doc, '')


def _render_choice(doc: Document, prepared: dict):
    payload = prepared['payload']
    options_count = len(payload['options'])
    for idx, option in enumerate(payload['options'], start=1):
        _add_text(doc, _list_item_text(idx, option['text'], idx == options_count))
    _add_text(doc, '')
    _add_text(doc, prepared['answer_instruction'])
    _add_text(doc, '')
    answer_table = doc.add_table(rows=1, cols=1)
    answer_table.style = 'Table Grid'
    _set_table_col_widths(answer_table, [3.0])
    _set_cell_text(answer_table.rows[0].cells[0], '')
    _add_text(doc, '')


def _render_open(doc: Document, prepared: dict):
    _add_text(doc, '')
    _add_text(doc, prepared['answer_instruction'])
    _add_text(doc, '')
    answer_table = doc.add_table(rows=1, cols=1)
    answer_table.style = 'Table Grid'
    _set_table_col_widths(answer_table, [6.0])
    _set_cell_text(answer_table.rows[0].cells[0], '')
    _add_text(doc, '')


def _render_task_block(doc: Document, prepared: dict):
    _add_text(doc, f'Задание {prepared["number"]}')
    _add_type_line(doc, prepared['type_name'])
    _add_text(doc, '')
    _add_italic_line(doc, 'Текст задания:')
    _add_text(doc, '')
    if prepared['task_intro']:
        _add_italic_line(doc, prepared['task_intro'])
        _add_text(doc, '')
    if prepared['prompt_text']:
        _add_text(doc, prepared['prompt_text'], align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    type_code = prepared['type_code']
    if type_code == TYPE_MATCHING:
        _render_matching(doc, prepared)
    elif type_code == TYPE_SEQUENCE:
        _render_sequence(doc, prepared)
    elif type_code == TYPE_MULTIPLE:
        _render_choice(doc, prepared)
    elif type_code == TYPE_SINGLE:
        _render_choice(doc, prepared)
    elif type_code == TYPE_OPEN:
        _render_open(doc, prepared)


def _add_specification_table(doc: Document, prepared_items: list[dict]):
    table = doc.add_table(rows=len(prepared_items) + 1, cols=4)
    table.style = 'Table Grid'
    _set_table_col_widths(table, [4.2, 5.4, 1.8, 4.0])
    headers = [
        'Код и наименование компетенции',
        'Код и наименование индикаторов сформированности компетенций',
        'Номер задания',
        'Тип задания',
    ]
    for idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[idx], header, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_idx, prepared in enumerate(prepared_items, start=1):
        _set_cell_text(table.rows[row_idx].cells[0], prepared['competence_text'])
        _set_cell_text(table.rows[row_idx].cells[1], prepared['indicator_text'])
        _set_cell_text(table.rows[row_idx].cells[2], str(prepared['number']))
        _set_cell_text(table.rows[row_idx].cells[3], prepared['type_name'])


def _add_keys_table(doc: Document, prepared_items: list[dict]):
    table = doc.add_table(rows=len(prepared_items) + 1, cols=4)
    table.style = 'Table Grid'
    _set_table_col_widths(table, [1.5, 4.8, 2.2, 6.9])
    headers = [
        'Номер задания',
        'Тип задания',
        'Верный ответ',
        'Критерии оценивания (баллы, получаемые за выполнение задания / характеристика правильности ответа)',
    ]
    for idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[idx], header, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_idx, prepared in enumerate(prepared_items, start=1):
        _set_cell_text(table.rows[row_idx].cells[0], str(prepared['number']))
        _set_cell_text(table.rows[row_idx].cells[1], prepared['type_name'])
        _set_cell_text(table.rows[row_idx].cells[2], prepared['answer_key'])
        _set_cell_text(table.rows[row_idx].cells[3], prepared['criteria'])


def _build_document(program_discipline, prepared_items: list[dict]) -> Document:
    if not MAKET_PATH.exists():
        logger.error('Word export template is missing', extra={'template_path': str(MAKET_PATH)})
        raise WordExportError('Не удалось сформировать Word-файл: не найден шаблон документа.')

    doc = Document(str(MAKET_PATH))
    _clear_document_content(doc)
    _configure_document_styles(doc)

    _add_text_styled(
        doc,
        'ОЦЕНОЧНЫЕ МАТЕРИАЛЫ',
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size_pt=TITLE_FONT_SIZE_PT,
    )
    _add_text_styled(
        doc,
        'по дисциплине',
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size_pt=TITLE_FONT_SIZE_PT,
    )
    _add_text_styled(
        doc,
        f'«{program_discipline.discipline.name}»',
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size_pt=TITLE_FONT_SIZE_PT,
    )
    _add_text(doc, '')

    _add_text_styled(doc, '1. Спецификация оценочных материалов', style='Normal', bold=True)
    _add_text(doc, '')
    _add_specification_table(doc, prepared_items)
    _add_text(doc, '')

    _add_text_styled(doc, '2. Тестовые задания', style='Normal', bold=True)
    _add_text(doc, '')
    for prepared in prepared_items:
        _render_task_block(doc, prepared)

    doc.add_page_break()
    _add_text_styled(doc, '3. Ключи и критерии оценивания', style='Normal', bold=True)
    _add_text(doc, '')
    _add_keys_table(doc, prepared_items)

    return doc


def generate_docx(program_id, discipline_id, filters):
    program_discipline = (
        ProgramDiscipline.objects.select_related('educational_program__program_profile', 'discipline')
        .filter(educational_program_id=program_id, discipline_id=discipline_id)
        .first()
    )
    if not program_discipline:
        raise WordExportNotFoundError('Связка программы и дисциплины не найдена.')

    items_queryset = _filtered_items(program_id, discipline_id, filters)
    items_count = items_queryset.count()
    if items_count > MAX_EXPORT_ITEMS:
        logger.info(
            'Word export item limit exceeded',
            extra={
                'program_id': program_id,
                'discipline_id': discipline_id,
                'items_count': items_count,
                'max_export_items': MAX_EXPORT_ITEMS,
            },
        )
        raise WordExportError(
            f'В экспорт попало слишком много заданий: {items_count}. '
            f'Уточните фильтры, максимум за один экспорт — {MAX_EXPORT_ITEMS}.'
        )

    items = list(items_queryset)
    if not items:
        raise WordExportNotFoundError('По выбранным фильтрам задания не найдены.')

    seed = f'{program_id}:{discipline_id}:{len(items)}:{datetime.now().strftime("%Y%m%d%H%M%S%f")}'
    rng = random.Random(seed)

    prepared_items = [
        _prepare_export_item(item, number=index, rng=rng)
        for index, item in enumerate(items, start=1)
    ]

    doc = _build_document(program_discipline, prepared_items)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()
