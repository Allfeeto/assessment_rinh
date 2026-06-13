from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from assessment.services import TYPE_MATCHING, TYPE_MULTIPLE, TYPE_OPEN, TYPE_SEQUENCE, TYPE_SINGLE

from .errors import WordExportError
from .preparers import _cleanup_text, _list_item_text, build_specification_groups

MAKET_PATH = Path(__file__).resolve().parent.parent / 'templates' / 'export' / 'maket.docx'
FONT_NAME = 'Times New Roman'
FONT_SIZE_PT = 12
TITLE_FONT_SIZE_PT = 14
logger = logging.getLogger(__name__)


def _set_cell_text(cell, text: str, *, bold: bool = False, align=None):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    _normalize_paragraph(paragraph)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    if bold:
        run.bold = True


def _clear_document_content(doc: Document):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith('sectPr'):
            continue
        body.remove(child)


def _normalize_paragraph(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)


def _configure_document_styles(doc: Document):
    for style_name in ('Normal', 'List Paragraph'):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = Pt(FONT_SIZE_PT)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1


def _set_table_col_widths(table, widths_cm: list[float]):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def _add_text(
    doc: Document,
    text: str = '',
    style: str = 'Normal',
    *,
    align=None,
    size_pt: int = FONT_SIZE_PT,
):
    paragraph = doc.add_paragraph(style=style)
    _normalize_paragraph(paragraph)
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(size_pt)
    return paragraph


def _add_text_styled(
    doc: Document,
    text: str = '',
    style: str = 'Normal',
    *,
    italic: bool = False,
    bold: bool = False,
    align=None,
    size_pt: int = FONT_SIZE_PT,
):
    paragraph = doc.add_paragraph(style=style)
    _normalize_paragraph(paragraph)
    if text:
        run = paragraph.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(size_pt)
        if italic:
            run.italic = True
        if bold:
            run.bold = True
    if align is not None:
        paragraph.alignment = align
    return paragraph


def _add_italic_line(doc: Document, text: str):
    paragraph = doc.add_paragraph(style='Normal')
    _normalize_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    run.italic = True
    return paragraph


def _add_type_line(doc: Document, type_name: str):
    paragraph = doc.add_paragraph(style='Normal')
    _normalize_paragraph(paragraph)
    label = paragraph.add_run('Тип задания: ')
    label.font.name = FONT_NAME
    label.font.size = Pt(FONT_SIZE_PT)
    label.italic = True
    value = paragraph.add_run(type_name)
    value.font.name = FONT_NAME
    value.font.size = Pt(FONT_SIZE_PT)
    return paragraph


def _render_matching(doc: Document, prepared: dict):
    item = prepared['item']
    payload = prepared['payload']

    _add_text(
        doc,
        'К каждой позиции, данной в левом столбце, подберите соответствующую позицию из правого столбца:',
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    _add_text(doc, '')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    left_header = table.rows[0].cells[0]
    _set_cell_text(
        left_header,
        _cleanup_text(item.left_column_title) or 'Левый столбец',
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    left_header.merge(table.rows[0].cells[1])
    right_header = table.rows[0].cells[2]
    _set_cell_text(
        right_header,
        _cleanup_text(item.right_column_title) or 'Правый столбец',
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    right_header.merge(table.rows[0].cells[3])

    right_entries = payload['right_items']
    left_pairs = payload['pairs']
    row_count = max(len(left_pairs), len(right_entries))

    for idx in range(row_count):
        cells = table.add_row().cells
        if idx < len(left_pairs):
            _set_cell_text(cells[0], payload['letters'][idx])
            _set_cell_text(cells[1], left_pairs[idx]['left'])
        if idx < len(right_entries):
            _set_cell_text(cells[2], str(idx + 1))
            _set_cell_text(cells[3], right_entries[idx]['text'])
    _set_table_col_widths(table, [0.8, 6.7, 0.8, 6.7])

    _add_text(doc, '')
    _add_text(doc, prepared['answer_instruction'])
    _add_text(doc, '')

    answer_table = doc.add_table(rows=2, cols=max(len(payload['letters']), 1))
    answer_table.style = 'Table Grid'
    _set_table_col_widths(answer_table, [1.0] * max(len(payload['letters']), 1))
    for idx, letter in enumerate(payload['letters'] or ['']):
        _set_cell_text(answer_table.rows[0].cells[idx], letter)
    for idx in range(max(len(payload['letters']), 1)):
        _set_cell_text(answer_table.rows[1].cells[idx], '')
    _add_text(doc, '')


def _render_sequence(doc: Document, prepared: dict):
    payload = prepared['payload']
    options_count = len(payload['steps'])
    for idx, step in enumerate(payload['steps'], start=1):
        _add_text(doc, _list_item_text(idx, step['text'], idx == options_count))
    _add_text(doc, '')
    _add_text(doc, prepared['answer_instruction'])
    _add_text(doc, '')
    answer_table = doc.add_table(rows=1, cols=payload['answer_cells'])
    answer_table.style = 'Table Grid'
    _set_table_col_widths(answer_table, [1.0] * payload['answer_cells'])
    for idx in range(payload['answer_cells']):
        _set_cell_text(answer_table.rows[0].cells[idx], '')
    _add_text(doc, '')


def _render_choice(doc: Document, prepared: dict):
    payload = prepared['payload']
    options_count = len(payload['options'])
    for idx, option in enumerate(payload['options'], start=1):
        _add_text(doc, _list_item_text(idx, option['text'], idx == options_count))
    _add_text(doc, '')
    _add_text(doc, prepared['answer_instruction'])
    _add_text(doc, '')
    answer_table = doc.add_table(rows=1, cols=1)
    answer_table.style = 'Table Grid'
    _set_table_col_widths(answer_table, [3.0])
    _set_cell_text(answer_table.rows[0].cells[0], '')
    _add_text(doc, '')


def _render_open(doc: Document, prepared: dict):
    _add_text(doc, '')
    _add_text(doc, prepared['answer_instruction'])
    _add_text(doc, '')
    answer_table = doc.add_table(rows=1, cols=1)
    answer_table.style = 'Table Grid'
    _set_table_col_widths(answer_table, [6.0])
    _set_cell_text(answer_table.rows[0].cells[0], '')
    _add_text(doc, '')


def _render_task_block(doc: Document, prepared: dict):
    type_code = prepared['type_code']
    _add_text(doc, f'Задание {prepared["number"]}')
    _add_type_line(doc, prepared['type_name'])
    _add_text(doc, '')
    _add_italic_line(doc, 'Текст задания:')
    _add_text(doc, '')
    if prepared['task_intro']:
        _add_text(doc, prepared['task_intro'])
        _add_text(doc, '')
    if prepared['prompt_text']:
        prompt_alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
            if type_code == TYPE_OPEN
            else WD_ALIGN_PARAGRAPH.JUSTIFY
        )
        _add_text(doc, prepared['prompt_text'], align=prompt_alignment)

    if type_code == TYPE_MATCHING:
        _render_matching(doc, prepared)
    elif type_code == TYPE_SEQUENCE:
        _render_sequence(doc, prepared)
    elif type_code == TYPE_MULTIPLE:
        _render_choice(doc, prepared)
    elif type_code == TYPE_SINGLE:
        _render_choice(doc, prepared)
    elif type_code == TYPE_OPEN:
        _render_open(doc, prepared)


def _add_specification_table(doc: Document, prepared_items: list[dict]):
    groups = build_specification_groups(prepared_items)
    table = doc.add_table(rows=len(groups) + 1, cols=4)
    table.style = 'Table Grid'
    _set_table_col_widths(table, [3.9, 6.9, 1.9, 3.5])
    headers = [
        'Код и наименование компетенции',
        'Код и наименование индикаторов сформированности компетенций',
        'Номера заданий',
        'Тип задания',
    ]
    for idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[idx], header, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_idx, group in enumerate(groups, start=1):
        _set_cell_text(table.rows[row_idx].cells[0], group['competence_text'])
        _set_cell_text(table.rows[row_idx].cells[1], group['indicator_text'])
        _set_cell_text(table.rows[row_idx].cells[2], group['numbers_text'])
        _set_cell_text(table.rows[row_idx].cells[3], group['type_name'])


def _add_keys_table(doc: Document, prepared_items: list[dict]):
    table = doc.add_table(rows=len(prepared_items) + 1, cols=4)
    table.style = 'Table Grid'
    _set_table_col_widths(table, [1.5, 4.8, 2.2, 6.9])
    headers = [
        'Номер задания',
        'Тип задания',
        'Верный ответ',
        'Критерии оценивания (баллы, получаемые за выполнение задания / характеристика правильности ответа)',
    ]
    for idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[idx], header, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_idx, prepared in enumerate(prepared_items, start=1):
        _set_cell_text(table.rows[row_idx].cells[0], str(prepared['number']))
        _set_cell_text(table.rows[row_idx].cells[1], prepared['type_name'])
        _set_cell_text(table.rows[row_idx].cells[2], prepared['answer_key'])
        _set_cell_text(table.rows[row_idx].cells[3], prepared['criteria'])


def build_document(program_discipline, prepared_items: list[dict]) -> Document:
    if not MAKET_PATH.exists():
        logger.error('Word export template is missing', extra={'template_path': str(MAKET_PATH)})
        raise WordExportError('Не удалось сформировать Word-файл: не найден шаблон документа.')

    doc = Document(str(MAKET_PATH))
    _clear_document_content(doc)
    _configure_document_styles(doc)
    discipline_name = getattr(program_discipline.discipline, 'name', '')

    _add_text_styled(
        doc,
        'ОЦЕНОЧНЫЕ МАТЕРИАЛЫ',
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size_pt=TITLE_FONT_SIZE_PT,
    )
    _add_text_styled(
        doc,
        'по дисциплине',
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size_pt=TITLE_FONT_SIZE_PT,
    )
    _add_text_styled(
        doc,
        f'«{discipline_name}»',
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size_pt=TITLE_FONT_SIZE_PT,
    )
    _add_text(doc, '')

    _add_text_styled(doc, '1. Спецификация оценочных материалов', style='Normal', bold=True)
    _add_text(doc, '')
    _add_specification_table(doc, prepared_items)
    _add_text(doc, '')

    _add_text_styled(doc, '2. Тестовые задания', style='Normal', bold=True)
    _add_text(doc, '')
    for prepared in prepared_items:
        _render_task_block(doc, prepared)

    doc.add_page_break()
    _add_text_styled(doc, '3. Ключи и критерии оценивания', style='Normal', bold=True)
    _add_text(doc, '')
    _add_keys_table(doc, prepared_items)

    return doc
